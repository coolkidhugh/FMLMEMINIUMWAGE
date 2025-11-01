import streamlit as st
import pandas as pd
import re
import io
import traceback
import json
from PIL import Image
from datetime import date, timedelta
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --- 新增 V3 依赖 ---
import base64
import asyncio
import requests # 假设 'requests' 库在环境中可用

# --- 移除 V2 依赖 (阿里云) ---
# try:
#     from alibabacloud_ocr_api20210707.client import Client as OcrClient
#     from alibabacloud_tea_openapi import models as open_api_models
#     # 导入 Table Recognize (表格识别) 模型
#     from alibabacloud_ocr_api20210707 import models as ocr_models
#     ALIYUN_SDK_AVAILABLE = True
# except ImportError:
#     ALIYUN_SDK_AVAILABLE = False
ALIYUN_SDK_AVAILABLE = False # 保留变量，但设为False

# ==============================================================================
# --- [核心功能 V3: Gemini Vision API] ---
# ==============================================================================

async def get_gemini_vision_analysis(image: Image.Image) -> dict:
    """
    调用 Gemini API (gemini-2.5-flash-preview-09-2025) 来分析图片并返回结构化JSON。
    """
    st.write("正在调用 Gemini Vision API...")
    
    # 1. 将Pillow Image转为Base64
    buffered = io.BytesIO()
    if image.mode == 'RGBA':
        image = image.convert('RGB')
    image.save(buffered, format="JPEG")
    base64_image_data = base64.b64encode(buffered.getvalue()).decode('utf-8')
    
    # 2. 定义API URL和Key (Key为空，由Canvas提供)
    apiKey = "" # 由Canvas环境自动提供
    apiUrl = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={apiKey}"

    # 3. 定义请求Prompt
    prompt = """
    请分析这张每日出租率对照表的图片。
    图片包含两个表格：'金陵楼' 和 '亚太商务楼'。
    
    你的任务是：
    1.  为这两个表格分别提取7天的数据。
    2.  如果一个单元格同时有打印数字和手写数字（通常是划掉了打印数字），请**只使用手写数字**，因为手写的是最终的正确值。
    3.  请提取以下列的数据：'日期', '星期', '当日预计', '当日实际', '周一预计', '平均房价'。
    4.  请忽略 '当日增加率', '当日实际' (周一的), '增加百分率' 这几列。
    5.  将所有提取的值作为字符串返回，保留原始格式（例如 '83.4%' 或 '5774'）。
    
    请严格按照下面提供的JSON schema格式返回数据。
    """

    # 4. 定义JSON Schema
    response_schema = {
      "type": "OBJECT",
      "properties": {
        "jinling": {
          "description": "金陵楼的7天数据",
          "type": "ARRAY",
          "items": {
            "type": "OBJECT",
            "properties": {
              "date": { "type": "STRING", "description": "日期, e.g., '20/10'" },
              "weekday": { "type": "STRING", "description": "星期, e.g., '一'" },
              "expected_today_raw": { "type": "STRING", "description": "当日预计, e.g., '78.4%'" },
              "actual_today_raw": { "type": "STRING", "description": "当日实际 (使用手写值), e.g., '83.4%'" },
              "expected_monday_raw": { "type": "STRING", "description": "周一预计, e.g., '83.4%'" },
              "avg_price_raw": { "type": "STRING", "description": "平均房价, e.g., '5774'" }
            },
            "required": ["date", "weekday", "expected_today_raw", "actual_today_raw", "expected_monday_raw", "avg_price_raw"]
          }
        },
        "yatai": {
          "description": "亚太商务楼的7天数据",
          "type": "ARRAY",
          "items": {
            "type": "OBJECT",
            "properties": {
              "date": { "type": "STRING", "description": "日期, e.g., '20/10'" },
              "weekday": { "type": "STRING", "description": "星期, e.g., '一'" },
              "expected_today_raw": { "type": "STRING", "description": "当日预计, e.g., '67.4%'" },
              "actual_today_raw": { "type": "STRING", "description": "当日实际 (使用手写值), e.g., '81.4%'" },
              "expected_monday_raw": { "type": "STRING", "description": "周一预计, e.g., '81.4%'" },
              "avg_price_raw": { "type": "STRING", "description": "平均房价, e.g., '7069'" }
            },
            "required": ["date", "weekday", "expected_today_raw", "actual_today_raw", "expected_monday_raw", "avg_price_raw"]
          }
        }
      },
      "required": ["jinling", "yatai"]
    }

    # 5. 构建Payload
    payload = {
        "contents": [
            {
                "role": "user",
                "parts": [
                    { "text": prompt },
                    {
                        "inlineData": {
                            "mimeType": "image/jpeg",
                            "data": base64_image_data
                        }
                    }
                ]
            }
        ],
        "generationConfig": {
            "responseMimeType": "application/json",
            "responseSchema": response_schema
        }
    }

    # 6. 发起API请求 (带指数退避)
    max_retries = 5
    delay = 1
    for attempt in range(max_retries):
        try:
            # 在 Streamlit 中，我们需要在一个单独的线程中运行阻塞的 I/O 操作
            response = await asyncio.to_thread(
                lambda: requests.post(
                    apiUrl,
                    headers={'Content-Type': 'application/json'},
                    data=json.dumps(payload),
                    timeout=60  # 60秒超时
                )
            )
            
            if response.status_code == 200:
                result = response.json()
                if (
                    result.get("candidates") and
                    result["candidates"][0].get("content", {}).get("parts") and
                    result["candidates"][0]["content"]["parts"][0].get("text")
                ):
                    st.write("API 调用成功，正在解析JSON...")
                    json_text = result["candidates"][0]["content"]["parts"][0]["text"]
                    return json.loads(json_text)
                else:
                    st.error(f"Gemini API 返回了意外的结构: {result}")
                    return None
            elif response.status_code == 429 or response.status_code >= 500:
                # 触发了重试 (速率限制或服务器错误)
                st.warning(f"API 返回 {response.status_code}, 正在重试... (第 {attempt + 1}/{max_retries} 次)")
                await asyncio.sleep(delay)
                delay *= 2
            else:
                # 客户端错误, 不重试
                st.error(f"Gemini API 请求失败: {response.status_code}")
                st.error(f"错误详情: {response.text}")
                return None
        except requests.exceptions.Timeout:
            st.warning(f"API 请求超时, G正在重试... (第 {attempt + 1}/{max_retries} 次)")
            await asyncio.sleep(delay)
            delay *= 2
        except Exception as e:
            st.error(f"调用 Gemini API 失败: {e}")
            st.code(traceback.format_exc())
            return None
            
    st.error(f"API 请求在 {max_retries} 次重试后失败。")
    return None

def populate_dataframe_from_json(ocr_data: dict, building_name: str) -> pd.DataFrame:
    """
    V3: 使用 Gemini 返回的结构化 JSON 来填充 DataFrame。
    """
    
    # 1. 准备一个空的默认DataFrame
    today = date.today()
    days = [(today + timedelta(days=i)) for i in range(7)]
    weekdays_zh_map = ["一", "二", "三", "四", "五", "六", "日"]
    
    initial_data = {
        "日期": [d.strftime("%m/%d") for d in days],
        "星期": [weekdays_zh_map[d.weekday()] for d in days],
        "当日预计 (%)": ["0.0"] * 7,
        "当日实际 (%)": ["0.0"] * 7,
        "周一预计 (%)": ["0.0"] * 7,
        "平均房价": ["0.0"] * 7
    }
    df = pd.DataFrame(initial_data)

    # 2. 检查
    if not ocr_data:
        st.warning("Gemini Vision 未返回数据。")
        return df

    building_key = "jinling" if building_name == "金陵楼" else "yatai"
    
    if building_key not in ocr_data:
        st.warning(f"在Gemini的JSON响应中未找到 '{building_key}' 键。")
        return df
        
    building_data = ocr_data[building_key]
    
    if not isinstance(building_data, list):
        st.warning(f"'{building_key}' 键对应的值不是一个列表。")
        return df

    # 3. 填充DataFrame
    for i in range(min(len(building_data), 7)): # 最多填7行
        row_data = building_data[i]
        try:
            # 我们可以选择覆盖日期和星期，或者只匹配
            # 覆盖更简单
            df.at[i, "日期"] = row_data.get("date", df.at[i, "日期"])
            df.at[i, "星期"] = row_data.get("weekday", df.at[i, "星期"])
            df.at[i, "当日预计 (%)"] = row_data.get("expected_today_raw", "0.0")
            df.at[i, "当日实际 (%)"] = row_data.get("actual_today_raw", "0.0")
            df.at[i, "周一预计 (%)"] = row_data.get("expected_monday_raw", "0.0")
            df.at[i, "平均房价"] = row_data.get("avg_price_raw", "0.0")
        except Exception as e:
            st.error(f"填充第 {i} 行数据时出错: {e}")
            continue
            
    st.write(f"成功从Gemini JSON填充了 '{building_name}' 的 {min(len(building_data), 7)} 行数据。")
    return df

# ==============================================================================
# --- [核心功能 V1: 计算和Word生成] ---
# (这部分代码和V1完全一样，因为它们只依赖DataFrame)
# ==============================================================================

def get_calc_value(value_str):
    """
    辅助函数：从 'xx.x%/yy.y%' 或 'xx.x' 字符串中提取用于计算的最后一个浮点数。
    """
    if isinstance(value_str, (int, float)):
        return value_str
    
    target_str = str(value_str)
    if '/' in target_str:
        target_str = target_str.split('/')[-1] # 取 '/' 后的部分
        
    # 清理常见的非数字字符
    target_str = target_str.replace('%', '').replace('i', '').strip()
        
    match = re.search(r'(\d+\.\d+)|(\d+)', target_str)
    if match:
        try:
            return float(match.group(1) or match.group(2))
        except (ValueError, TypeError):
            return 0.0
    return 0.0

def calculate_rates(df_in):
    """
    使用提取的浮点数计算“当日增加率”和“增加百分率”。
    """
    df = df_in.copy()
    
    try:
        # 1. 提取用于计算的浮点数列
        calc_expected = df["当日预计 (%)"].apply(get_calc_value)
        calc_actual = df["当日实际 (%)"].apply(get_calc_value)
        calc_monday_expected = df["周一预计 (%)"].apply(get_calc_value)
        
        # 2. 计算增加率
        df["当日增加率 (%)"] = calc_actual - calc_expected
        df["增加百分率 (%)"] = calc_actual - calc_monday_expected
        
        # 3. 格式化为带正负号的字符串
        df["当日增加率 (%)"] = df["当日增加率 (%)"].apply(lambda x: f"{x:+.1f}%")
        df["增加百分率 (%)"] = df["增加百分率 (%)"].apply(lambda x: f"{x:+.1f}%")
        
        # 4. 计算本周总结
        actual_sum = calc_actual.sum()
        monday_sum = calc_monday_expected.sum()
        
        summary = {
            "本周实际": f"{actual_sum:.1f}%",
            "周一预测": f"{monday_sum:.1f}%",
            "实际增加": f"{(actual_sum - monday_sum):+.1f}%" # 使用带符号的格式
        }
        
        return df, summary
    except Exception as e:
        st.error(f"计算比率时出错: {e}")
        st.code(traceback.format_exc())
        return df_in, {} # 返回原始表和空总结

def create_word_doc(jl_df, yt_df, jl_summary, yt_summary):
    """
    将两个DataFrame和它们的总结数据生成一个Word文档。
    """
    try:
        doc = Document()
        # 设置中文字体
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体' )
        doc.styles['Normal'].font.size = Pt(10)
        
        # --- 金陵楼 ---
        doc.add_heading("每日出租率对照表", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("金陵楼", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 调整列顺序以匹配原始表格
        jl_cols_ordered = ["日期", "星期", "当日预计 (%)", "当日实际 (%)", "当日增加率 (%)", "周一预计 (%)", "增加百分率 (%)", "平均房价"]
        jl_df_ordered = jl_df[jl_cols_ordered]

        table_jl = doc.add_table(rows=jl_df_ordered.shape[0] + 1, cols=jl_df_ordered.shape[1])
        table_jl.style = 'Table Grid'
        
        # 写表头
        hdr_cells_jl = table_jl.rows[0].cells
        for i, col_name in enumerate(jl_df_ordered.columns):
            hdr_cells_jl[i].text = col_name
            hdr_cells_jl[i].paragraphs[0].runs[0].font.bold = True

        # 写数据
        for index, row in jl_df_ordered.iterrows():
            row_cells = table_jl.rows[index + 1].cells
            for i, col_name in enumerate(jl_df_ordered.columns):
                row_cells[i].text = str(row[col_name])
        
        # 写金陵楼总结
        doc.add_paragraph(
            f"本周实际： {jl_summary.get('本周实际', 'N/A')}     "
            f"周一预测： {jl_summary.get('周一预测', 'N/A')}     "
            f"实际增加： {jl_summary.get('实际增加', 'N/A')}"
        )

        # --- 亚太商务楼 ---
        doc.add_paragraph("亚太商务楼", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        yt_cols_ordered = ["日期", "星期", "当日预计 (%)", "当日实际 (%)", "当日增加率 (%)", "周一预计 (%)", "增加百分率 (%)", "平均房价"]
        yt_df_ordered = yt_df[yt_cols_ordered]

        table_yt = doc.add_table(rows=yt_df_ordered.shape[0] + 1, cols=yt_df_ordered.shape[1])
        table_yt.style = 'Table Grid'
        
        hdr_cells_yt = table_yt.rows[0].cells
        for i, col_name in enumerate(yt_df_ordered.columns):
            hdr_cells_yt[i].text = col_name
            hdr_cells_yt[i].paragraphs[0].runs[0].font.bold = True

        for index, row in yt_df_ordered.iterrows():
            row_cells = table_yt.rows[index + 1].cells
            for i, col_name in enumerate(yt_df_ordered.columns):
                row_cells[i].text = str(row[col_name])
        
        doc.add_paragraph(
            f"本周实际： {yt_summary.get('本周实际', 'N/A')}     "
            f"周一预测： {yt_summary.get('周一预测', 'N/A')}     "
            f"实际增加： {yt_summary.get('实际增加', 'N/A')}"
        )

        # 设置列宽
        widths = [Inches(0.6), Inches(0.5), Inches(0.8), Inches(1.0), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8)]
        for table in [table_jl, table_yt]:
            for i, width in enumerate(widths):
                if i < len(table.columns):
                    for cell in table.columns[i].cells:
                        cell.width = width
        
        # 存到内存
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f.getvalue()

    except Exception as e:
        st.error(f"生成Word文档时出错: {e}")
        st.code(traceback.format_exc())
        return None

# ==============================================================================
# --- [Streamlit 界面 V3] ---
# ==============================================================================
def run_ocr_calculator_app():
    st.title("OCR出租率计算器 (V3 - Gemini Vision)")
    st.markdown("1. 上传手写表格的照片。")
    st.markdown("2. 使用 **Gemini Vision API** 智能解析表格（优先读取手写值）。")
    st.markdown("3. **人工核对**下方的可编辑表格，修正识别错误的数字。")
    st.markdown("4. 点击“计算”按钮，生成最终报表并下载Word文档。")

    uploaded_file = st.file_uploader("上传图片文件", type=["png", "jpg", "jpeg", "bmp"], key="ocr_calc_uploader_v3") # 新Key

    if uploaded_file:
        image = Image.open(uploaded_file)
        st.image(image, caption="上传的图片", width=300)

        if st.button("开始识别 (Gemini Vision)", type="primary"): # 新按钮
            with st.spinner('正在调用 Gemini Vision API (这可能需要一些时间)...'):
                
                # 在 Streamlit 中运行异步函数
                try:
                    ocr_data = asyncio.run(get_gemini_vision_analysis(image))
                except Exception as e:
                    st.error(f"运行异步Gemini任务时出错: {e}")
                    ocr_data = None
                
                if ocr_data:
                    st.session_state.ocr_data = ocr_data
                    st.info("API调用成功，正在解析返回的JSON...")
                    # *** V3: 使用新的解析函数 ***
                    st.session_state.jl_df = populate_dataframe_from_json(ocr_data, "金陵楼")
                    st.session_state.yt_df = populate_dataframe_from_json(ocr_data, "亚太商务楼")
                    st.success("解析完成！请检查下面的表格，手动修正错误。")
                    
                    with st.expander("查看 Gemini 返回的原始JSON数据"):
                        st.json(ocr_data)
                else:
                    st.error("Gemini API 未返回有效数据。")
                    st.session_state.jl_df = populate_dataframe_from_json(None, "金陵楼") # 生成空表
                    st.session_state.yt_df = populate_dataframe_from_json(None, "亚太商务楼")
    
    # --- 表格编辑区 ---
    if 'jl_df' in st.session_state:
        st.divider()
        st.subheader("金陵楼 (请在此处编辑)")
        st.session_state.jl_df_edited = st.data_editor(
            st.session_state.jl_df,
            column_config={
                "当日预计 (%)": st.column_config.TextColumn(label="当日预计 (%)", width="small"),
                "当日实际 (%)": st.column_config.TextColumn(label="当日实际 (%)", width="medium"),
                "周一预计 (%)": st.column_config.TextColumn(label="周一预计 (%)", width="small"),
                "平均房价": st.column_config.TextColumn(label="平均房价", width="small"),
                "日期": st.column_config.TextColumn(label="日期", disabled=True),
                "星期": st.column_config.TextColumn(label="星期", disabled=True),
            },
            num_rows="fixed",
            key="editor_jl_v3" # 新Key
        )
        
        st.subheader("亚太商务楼 (请在此处编辑)")
        st.session_state.yt_df_edited = st.data_editor(
            st.session_state.yt_df,
            column_config={
                "当日预计 (%)": st.column_config.TextColumn(label="当日预计 (%)", width="small"),
                "当日实际 (%)": st.column_config.TextColumn(label="当日实际 (%)", width="medium"),
                "周一预计 (%)": st.column_config.TextColumn(label="周一预计 (%)", width="small"),
                "平均房价": st.column_config.TextColumn(label="平均房价", width="small"),
                "日期": st.column_config.TextColumn(label="日期", disabled=True),
                "星期": st.column_config.TextColumn(label="星期", disabled=True),
            },
            num_rows="fixed",
            key="editor_yt_v3" # 新Key
        )

        if st.button("重新计算最终结果", type="primary"):
            # 使用编辑后的数据进行计算
            jl_df_final, jl_summary = calculate_rates(st.session_state.jl_df_edited)
            yt_df_final, yt_summary = calculate_rates(st.session_state.yt_df_edited)
            
            st.session_state.jl_df_final = jl_df_final
            st.session_state.yt_df_final = yt_df_final
            st.session_state.jl_summary = jl_summary
            st.session_state.yt_summary = yt_summary
            
            st.balloons()
            st.success("计算完成！")

    # --- 最终结果展示 ---
    if 'jl_df_final' in st.session_state:
        st.divider()
        st.subheader("最终结果 (金陵楼)")
        
        # 辅助函数，用于安全格式化
        def format_price(x):
            val = get_calc_value(x)
            return f"{val:.2f}" if val is not None else "0.00"

        st.dataframe(st.session_state.jl_df_final.style.format({
           "平均房价": format_price
        }))
        st.markdown(
            f"**本周实际：** `{st.session_state.jl_summary.get('本周实际', 'N/A')}` | "
            f"**周一预测：** `{st.session_state.jl_summary.get('周一预测', 'N/A')}` | "
            f"**实际增加：** `{st.session_state.jl_summary.get('实际增加', 'N/A')}`"
        )
        
        st.subheader("最终结果 (亚太商务楼)")
        st.dataframe(st.session_state.yt_df_final.style.format({
            "平均房价": format_price
        }))
        st.markdown(
            f"**本周实际：** `{st.session_state.yt_summary.get('本周实际', 'N/A')}` | "
            f"**周一预测：** `{st.session_state.yt_summary.get('周一预测', 'N/A')}` | "
            f"**实际增加：** `{st.session_state.yt_summary.get('实际增加', 'N/A')}`"
        )
        
        # 生成Word
        doc_data = create_word_doc(
            st.session_state.jl_df_final, 
            st.session_state.yt_df_final, 
            st.session_state.jl_summary, 
            st.session_state.yt_summary
        )
        
        if doc_data:
            st.download_button(
                label="下载Word文档",
                data=doc_data,
                file_name="每日出租率对照表_V3_Gemini.docx", # 新文件名
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# --- 主程序入口 ---
if __name__ == "__main__":
    # 设置页面标题
    st.set_page_config(page_title="OCR出租率计算器 V3")
    run_ocr_calculator_app()

